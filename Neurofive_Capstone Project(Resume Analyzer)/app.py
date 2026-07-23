"""
ResumeForge AI — Flask backend
================================
Powers the index.html single-page frontend:
  - Serves the UI
  - Extracts text from uploaded resumes (.pdf, .docx, .txt)
  - Talks to an LLM through the BlueMinds AI API (OpenAI-compatible) to:
      * score / analyze a resume against a target company + job description
      * generate a tailored cover letter
  - Exposes a small "company profile" catalogue used to steer the AI's
    company-tailored recommendations

See README.md for setup instructions.
"""

import json
import os
import re

from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request
from openai import OpenAI
from pypdf import PdfReader
from docx import Document
from werkzeug.utils import secure_filename

load_dotenv()

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

BLUESMINDS_API_KEY = os.getenv("BLUESMINDS_API_KEY", "").strip()
BLUESMINDS_BASE_URL = os.getenv("BLUESMINDS_BASE_URL", "https://api.bluesminds.com/v1").strip()
BLUESMINDS_MODEL = os.getenv("BLUESMINDS_MODEL", "gpt-4o-mini").strip()

MAX_UPLOAD_MB = 8
ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_MB * 1024 * 1024

# OpenAI-compatible client pointed at the BlueMinds gateway.
# BlueMinds (https://api.bluesminds.com) exposes an OpenAI-compatible
# /v1/chat/completions endpoint, so the official `openai` SDK works as-is —
# we just swap the base_url and use a BlueMinds API key.
client = None
if BLUESMINDS_API_KEY:
    client = OpenAI(api_key=BLUESMINDS_API_KEY, base_url=BLUESMINDS_BASE_URL)


# ---------------------------------------------------------------------------
# Company profiles — used to steer "company-tailored" recommendations.
# Feel free to extend this dict with more companies.
# ---------------------------------------------------------------------------

COMPANIES = {
    "google": {
        "name": "Google",
        "profile": (
            "Data-driven engineering culture, obsessed with scale, ambiguity "
            "and structured problem solving. Values Googleyness (collaboration, "
            "intellectual humility), quantified impact statements, and clear "
            "signal on CS fundamentals / systems thinking."
        ),
    },
    "amazon": {
        "name": "Amazon",
        "profile": (
            "Leadership Principles are everything (Customer Obsession, "
            "Ownership, Bias for Action, Dive Deep, Deliver Results). Resumes "
            "and interviews are evaluated through STAR-style ownership stories "
            "with measurable business outcomes."
        ),
    },
    "microsoft": {
        "name": "Microsoft",
        "profile": (
            "Growth-mindset culture, cross-team collaboration, and customer "
            "impact across a broad product portfolio (Azure, M365, Copilot). "
            "Values learning agility and inclusive teamwork alongside technical depth."
        ),
    },
    "meta": {
        "name": "Meta",
        "profile": (
            "Move-fast, high-ownership engineering culture focused on "
            "shipping velocity, bold bets, and measurable product/user impact "
            "at massive scale."
        ),
    },
    "apple": {
        "name": "Apple",
        "profile": (
            "Design-obsessed, secrecy-conscious, craftsmanship-first culture. "
            "Rewards precision, attention to detail, and cross-functional "
            "collaboration between engineering and design."
        ),
    },
    "netflix": {
        "name": "Netflix",
        "profile": (
            "Freedom & Responsibility culture — highly autonomous senior "
            "operators who make judgment calls with minimal process. Values "
            "candor, high performance, and business-outcome-driven storytelling."
        ),
    },
    "startup": {
        "name": "Early-Stage Startup",
        "profile": (
            "Generalist, scrappy, high-ownership culture with limited "
            "resources. Values breadth of skills, bias for action, and "
            "evidence of shipping things end-to-end with little oversight."
        ),
    },
    "generic": {
        "name": "Generic / Other",
        "profile": (
            "No specific company target — evaluate against strong general "
            "industry best practices for resumes and interviews."
        ),
    },
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_stream) -> str:
    reader = PdfReader(file_stream)
    text_parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_stream) -> str:
    document = Document(file_stream)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def extract_text_from_txt(file_stream) -> str:
    raw = file_stream.read()
    try:
        return raw.decode("utf-8").strip()
    except UnicodeDecodeError:
        return raw.decode("latin-1", errors="ignore").strip()


def extract_json(raw_text: str) -> dict:
    """Best-effort extraction of a JSON object from an LLM response,
    in case the model wraps it in markdown fences or adds stray text."""
    raw_text = raw_text.strip()
    raw_text = re.sub(r"^```(json)?", "", raw_text.strip(), flags=re.IGNORECASE).strip()
    raw_text = re.sub(r"```$", "", raw_text.strip()).strip()
    try:
        return json.loads(raw_text)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", raw_text, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        raise


def call_ai(system_prompt: str, user_prompt: str) -> dict:
    """Calls the BlueMinds chat completion endpoint and returns parsed JSON."""
    if client is None:
        raise RuntimeError("BlueMinds API key is not configured on the server.")

    response = client.chat.completions.create(
        model=BLUESMINDS_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.6,
        response_format={"type": "json_object"},
    )
    content = response.choices[0].message.content
    return extract_json(content)


def get_company_context(company_id: str) -> dict:
    return COMPANIES.get(company_id, COMPANIES["generic"])


# ---------------------------------------------------------------------------
# Prompt builders
# ---------------------------------------------------------------------------

ANALYZE_SYSTEM_PROMPT = """You are ResumeForge AI, a brutally honest, world-class resume
reviewer, ATS optimization expert, and technical interview coach. You analyze resumes
against a target company's culture and an (optional) job description, then return a
single, strictly valid JSON object — no markdown, no commentary, no code fences.

The JSON object MUST follow exactly this schema (all fields required):

{
  "overall_score": <integer 0-100>,
  "score_breakdown": {
    "content_quality": <integer 0-100>,
    "formatting": <integer 0-100>,
    "keyword_optimization": <integer 0-100>,
    "impact_metrics": <integer 0-100>,
    "company_fit": <integer 0-100>
  },
  "executive_summary": "<2-4 sentence honest summary>",
  "strengths": ["<short strength>", ...],
  "weaknesses": ["<short weakness>", ...],
  "missing_keywords": ["<keyword>", ...],
  "suggested_improvements": [
    {"section": "<resume section>", "current": "<verbatim-ish snippet or paraphrase>",
     "suggested": "<rewritten stronger version>", "impact": "<why this helps>"}
  ],
  "tailored_recommendations": ["<company-specific recommendation>", ...],
  "interview_prep": {
    "likely_questions": [
      {"type": "<Behavioral|Technical|Situational|Culture Fit>",
       "question": "<question text>", "preparation_tip": "<short tip>"}
    ],
    "talking_points": ["<point to emphasize>", ...],
    "red_flags_to_address": ["<gap or concern to proactively address>", ...]
  },
  "ats_optimization": {
    "score": <integer 0-100>,
    "issues": ["<ATS issue>", ...],
    "fixes": ["<concrete fix>", ...]
  },
  "action_plan": [
    {"priority": "<high|medium|low>", "action": "<concrete next step>",
     "time_estimate": "<e.g. 15 min>"}
  ]
}

Be specific and concrete. Use real numbers/metrics from the resume where possible.
Tailor everything to the target company's culture and the job description provided.
Return ONLY the JSON object."""


COVER_LETTER_SYSTEM_PROMPT = """You are ResumeForge AI's cover letter writer. Given a
resume, an optional job description, a target company, an experience level and a
desired tone, write a compelling, specific, non-generic cover letter (250-400 words)
that references real details from the resume. Return a single strictly valid JSON
object — no markdown, no commentary, no code fences — with exactly this schema:

{
  "subject_line": "<short compelling email subject line>",
  "cover_letter": "<the full cover letter text, ready to send>"
}

Return ONLY the JSON object."""


def build_analyze_user_prompt(resume_text, job_description, company, experience_level):
    return f"""TARGET COMPANY: {company['name']}
COMPANY CULTURE PROFILE: {company['profile']}
CANDIDATE EXPERIENCE LEVEL: {experience_level}

JOB DESCRIPTION (may be empty):
{job_description or '(not provided — evaluate generically for this role/level)'}

RESUME TEXT:
{resume_text}

Analyze this resume now and return the JSON object described in your instructions."""


def build_cover_letter_user_prompt(resume_text, job_description, company, experience_level, tone):
    return f"""TARGET COMPANY: {company['name']}
COMPANY CULTURE PROFILE: {company['profile']}
CANDIDATE EXPERIENCE LEVEL: {experience_level}
DESIRED TONE: {tone or 'Professional and confident'}

JOB DESCRIPTION (may be empty):
{job_description or '(not provided — write a strong general cover letter for this role/level)'}

RESUME TEXT:
{resume_text}

Write the cover letter now and return the JSON object described in your instructions."""


# ---------------------------------------------------------------------------
# Routes — frontend
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


# ---------------------------------------------------------------------------
# Routes — API
# ---------------------------------------------------------------------------

@app.route("/api/status")
def api_status():
    return jsonify({
        "api_configured": client is not None,
        "model": BLUESMINDS_MODEL if client is not None else None,
    })


@app.route("/api/companies")
def api_companies():
    companies = [
        {"id": cid, "name": c["name"]}
        for cid, c in COMPANIES.items()
        if cid != "generic"
    ]
    companies.sort(key=lambda c: c["name"])
    return jsonify(companies)


@app.route("/api/upload-resume", methods=["POST"])
def api_upload_resume():
    if "file" not in request.files:
        return jsonify({"error": True, "message": "No file was uploaded."}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": True, "message": "No file was selected."}), 400

    if not allowed_file(file.filename):
        return jsonify({
            "error": True,
            "message": "Unsupported file type. Please upload a .pdf, .docx, or .txt file.",
        }), 400

    filename = secure_filename(file.filename)
    ext = filename.rsplit(".", 1)[1].lower()

    try:
        if ext == "pdf":
            text = extract_text_from_pdf(file.stream)
        elif ext == "docx":
            text = extract_text_from_docx(file.stream)
        else:
            text = extract_text_from_txt(file.stream)
    except Exception as exc:  # noqa: BLE001
        return jsonify({
            "error": True,
            "message": f"Could not read that file: {exc}",
        }), 400

    if not text or len(text.strip()) < 20:
        return jsonify({
            "error": True,
            "message": "Couldn't extract readable text from that file. Try pasting your resume text instead.",
        }), 400

    return jsonify({"resume_text": text, "filename": filename})


@app.route("/api/analyze", methods=["POST"])
def api_analyze():
    payload = request.get_json(silent=True) or {}
    resume_text = (payload.get("resume_text") or "").strip()
    job_description = (payload.get("job_description") or "").strip()
    company_id = (payload.get("company_name") or "generic").strip()
    experience_level = (payload.get("experience_level") or "Mid-Level").strip()

    if len(resume_text) < 50:
        return jsonify({
            "error": True,
            "message": "Resume text must be at least 50 characters long.",
        }), 400

    if client is None:
        return jsonify({
            "error": True,
            "message": "AI engine is not configured. Set BLUESMINDS_API_KEY on the server.",
        }), 503

    company = get_company_context(company_id)

    try:
        result = call_ai(
            ANALYZE_SYSTEM_PROMPT,
            build_analyze_user_prompt(resume_text, job_description, company, experience_level),
        )
        return jsonify(result)
    except Exception as exc:  # noqa: BLE001
        return jsonify({
            "error": True,
            "message": f"Analysis failed: {exc}",
        }), 502


@app.route("/api/cover-letter", methods=["POST"])
def api_cover_letter():
    payload = request.get_json(silent=True) or {}
    resume_text = (payload.get("resume_text") or "").strip()
    job_description = (payload.get("job_description") or "").strip()
    company_id = (payload.get("company_name") or "generic").strip()
    experience_level = (payload.get("experience_level") or "Mid-Level").strip()
    tone = (payload.get("tone") or "Professional and confident").strip()

    if len(resume_text) < 50:
        return jsonify({
            "error": True,
            "message": "Resume text must be at least 50 characters long.",
        }), 400

    if client is None:
        return jsonify({
            "error": True,
            "message": "AI engine is not configured. Set BLUESMINDS_API_KEY on the server.",
        }), 503

    company = get_company_context(company_id)

    try:
        result = call_ai(
            COVER_LETTER_SYSTEM_PROMPT,
            build_cover_letter_user_prompt(resume_text, job_description, company, experience_level, tone),
        )
        return jsonify(result)
    except Exception as exc:  # noqa: BLE001
        return jsonify({
            "error": True,
            "message": f"Cover letter generation failed: {exc}",
        }), 502


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "true").lower() == "true"
    app.run(host="0.0.0.0", port=port, debug=debug)
